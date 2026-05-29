"""Document processor for text extraction from uploaded files.

Supports: .txt, .pdf, .docx, and image files (.png, .jpg, .jpeg).
Uses pure-Python libraries so no external system tools are needed.
Images are processed by EasyOCR for text extraction.
"""

import io
from typing import Optional
from fastapi import UploadFile, HTTPException

# PyPDF2 — pure Python PDF text extraction (no Poppler required)
from PyPDF2 import PdfReader

# python-docx — extract text from .docx Word documents
from docx import Document


async def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    """Extract text from a PDF file using PyPDF2.
    
    This uses direct text extraction, which works well for text-based PDFs
    (slides exported as PDF, Word docs saved as PDF, etc.).
    
    Args:
        pdf_bytes: Raw bytes of the PDF file.
        
    Returns:
        Extracted text string.
    """
    try:
        # Create a PDF reader from the byte stream
        reader = PdfReader(io.BytesIO(pdf_bytes))
        
        extracted_text = []
        # Loop through each page and extract text
        for idx, page in enumerate(reader.pages):
            page_text = page.extract_text() or ""
            if page_text.strip():
                extracted_text.append(f"--- Page {idx + 1} ---\n{page_text}")
        
        result = "\n\n".join(extracted_text).strip()
        
        if not result:
            raise HTTPException(
                status_code=400,
                detail="Could not extract text from this PDF. It may be a scanned/image-only PDF."
            )
        
        return result
    except HTTPException:
        raise  # Re-raise our own HTTP exceptions
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to process PDF: {str(e)}")


async def extract_text_from_docx(docx_bytes: bytes) -> str:
    """Extract text from a .docx Word document.
    
    Args:
        docx_bytes: Raw bytes of the DOCX file.
        
    Returns:
        Extracted text string.
    """
    try:
        # Load the DOCX document from bytes
        doc = Document(io.BytesIO(docx_bytes))
        
        # Extract text from all paragraphs
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        
        # Also extract text from tables (students often have notes in tables)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        
        result = "\n\n".join(paragraphs).strip()
        
        if not result:
            raise HTTPException(
                status_code=400,
                detail="Could not extract text from this DOCX. The document appears to be empty."
            )
        
        return result
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to process DOCX: {str(e)}")

async def extract_text_from_image(image_bytes: bytes) -> str:
    """Extract text from an image using EasyOCR.
    
    EasyOCR is a self-contained Python OCR library — no external tools
    like Tesseract or Poppler needed. It handles printed text well and
    has decent handwriting support.
    
    Args:
        image_bytes: Raw bytes of the image file.
        
    Returns:
        Extracted text string.
    """
    try:
        import easyocr
        import numpy as np
        from PIL import Image
        
        # Load and convert image to numpy array for EasyOCR
        img = Image.open(io.BytesIO(image_bytes))
        img_array = np.array(img)
        
        # Initialize EasyOCR reader (English by default, cached after first run)
        reader = easyocr.Reader(['en'], gpu=False)
        
        # Run OCR on the image
        results = reader.readtext(img_array, detail=0, paragraph=True)
        
        text = "\n".join(results).strip()
        
        if not text:
            raise HTTPException(
                status_code=400,
                detail="Could not find any readable text in this image. Please upload a clearer image of your notes."
            )
        
        return text
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to process image: {str(e)}")


async def process_file(file: UploadFile) -> str:
    """Read an uploaded file and route it to the appropriate processor based on its content type.
    
    Args:
        file: The FastAPI UploadFile object.
        
    Returns:
        Extracted text string.
    """
    # Read the file's contents into memory
    file_bytes = await file.read()
    content_type = file.content_type or ""
    filename = (file.filename or "").lower()

    # Route based on MIME type and file extension
    # PDF files
    if content_type == "application/pdf" or filename.endswith(".pdf"):
        return await extract_text_from_pdf(file_bytes)
    
    # DOCX files
    elif (content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
          or filename.endswith(".docx")):
        return await extract_text_from_docx(file_bytes)
    
    # Image files
    elif content_type.startswith("image/") or filename.endswith((".png", ".jpg", ".jpeg")):
        return await extract_text_from_image(file_bytes)
    
    # Plain text / JSON files
    elif content_type.startswith("text/") or content_type == "application/json" or filename.endswith(".txt"):
        return file_bytes.decode("utf-8").strip()
    
    else:
        # Attempt to decode as text as a fallback, otherwise raise an error
        try:
            return file_bytes.decode("utf-8").strip()
        except UnicodeDecodeError:
            raise HTTPException(
                status_code=400, 
                detail=f"Unsupported file type: {content_type}. Please upload a .txt, .pdf, .docx, or image file."
            )
